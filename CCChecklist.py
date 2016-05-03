#!/usr/bin/env python
# -*- coding: utf-8 -*-

__license__ = """

CCChecklist

Author: https://twitter.com/1_mod_m/

Project site: https://github.com/1modm

Copyright (c) 2016, MM
All rights reserved.

Redistribution and use in source and binary forms, with or without
modification, are permitted provided that the following conditions
are met:

1. Redistributions of source code must retain the above copyright
   notice, this list of conditions and the following disclaimer.
2. Redistributions in binary form must reproduce the above copyright
   notice, this list of conditions and the following disclaimer in the
   documentation and/or other materials provided with the distribution.
3. Neither the name of copyright holders nor the names of its
   contributors may be used to endorse or promote products derived
   from this software without specific prior written permission.

THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS
''AS IS'' AND ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED
TO, THE IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR
PURPOSE ARE DISCLAIMED.  IN NO EVENT SHALL COPYRIGHT HOLDERS OR CONTRIBUTORS
BE LIABLE FOR ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR
CONSEQUENTIAL DAMAGES (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF
SUBSTITUTE GOODS OR SERVICES; LOSS OF USE, DATA, OR PROFITS; OR BUSINESS
INTERRUPTION) HOWEVER CAUSED AND ON ANY THEORY OF LIABILITY, WHETHER IN
CONTRACT, STRICT LIABILITY, OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE)
ARISING IN ANY WAY OUT OF THE USE OF THIS SOFTWARE, EVEN IF ADVISED OF THE
POSSIBILITY OF SUCH DAMAGE.
"""

#------------------------------------------------------------------------------
# Modules
#------------------------------------------------------------------------------

import os
import sys
import hashlib
from datetime import datetime
import sqlite3
import shutil
from thirdparty.color.termcolor import colored
from lib.txtoutput import create_txt_file, print_title_txt, print_result_linesep_txt, print_result_txt


#------------------------------------------------------------------------------
# Python version check.
#------------------------------------------------------------------------------

if __name__ == "__main__":
    if sys.version_info < (2, 7) or sys.version_info >= (3, 0):
        show_banner()
        print ("[!] You must use Python version 2.7 or above")
        sys.exit(1)

try:
    import argparse
except ImportError:
    print 'argparse required. Install using `pip install argparse`'
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print 'docx required. Install using `pip install python-docx`'
    sys.exit(1)


try:
    import xlsxwriter
except ImportError:
    print 'xlsxwriter required. Install using `pip install xlsxwriter`'
    sys.exit(1)


#------------------------------------------------------------------------------
# Command line parser using argparse
#------------------------------------------------------------------------------

def cmdline_parser():
    parser = argparse.ArgumentParser(conflict_handler='resolve', add_help=True,
        description='Example: python %(prog)s -eal EAL4 -ac ALC_FLR.1 ATE_DPT.2 -p 1 data/ccdb.sqlite3', version='CChecklist 0.1',
        usage="python %(prog)s [OPTIONS]")

    # Mandatory
    parser.add_argument('database', action="store")

    # Optional
    parser.add_argument('-l', '--list', action='store_true', default=False, dest='list',
        help='List all Assurance Components')
    parser.add_argument('-eal', action='store', default=False, dest='EAL',
        help='EAL to use')
    parser.add_argument('-ac', metavar='AC', action='store', default=False, nargs='+', 
        help='Add Assurance Component')
    parser.add_argument('-p', action='store', default=False, dest='npara',
        help='Number of paragraphs to write in each workunit')
    parser.add_argument('-ask', action='store_true', default=False, dest='ask')

    return parser


###############################################################################
#------------------------------------------------------------------------------
# Class DatabaseManager
#------------------------------------------------------------------------------
###############################################################################

class DatabaseManager(object):
    def __init__(self, db):
        self.conn = sqlite3.connect(db)
        self.conn.execute('pragma foreign_keys = on')
        self.conn.commit()
        self.cur = self.conn.cursor()

    def query(self, arg):
        self.cur.execute(arg)
        self.conn.commit()
        return self.cur

    def __del__(self):
        self.conn.close()


###############################################################################
#------------------------------------------------------------------------------
# Class Workunit
#------------------------------------------------------------------------------
###############################################################################

class Workunit(object):

    def AugmentationList(self, filedb):

        if not (os.path.isfile(filedb)):
          print os.linesep + "Error loading database. Must run CCParser.py first" + os.linesep
          sys.exit(1)
        # Database
        try:
            # Create or open a SQLite3 DB
            dbmgr1 = DatabaseManager(filedb)
            dbmgr2 = DatabaseManager(filedb)
            dbmgr3 = DatabaseManager(filedb)
            dbmgr4 = DatabaseManager(filedb)
            dbmgr5 = DatabaseManager(filedb)
            dbmgr6 = DatabaseManager(filedb)
            dbmgr7 = DatabaseManager(filedb)
            dbmgr8 = DatabaseManager(filedb)
            dbmgr9 = DatabaseManager(filedb)
        except NameError:
            print os.linesep + "Error loading database" + os.linesep
            sys.exit(1)

        # Sub Activity
        query1 = "select child3.idelement, child3.paratext from child3 where child3.element=\"a-component\""

        print "Assurance Components:" + os.linesep * 2

        for row in dbmgr1.query(query1):
          print "Assurance Component: " + row[1]


    def queryworkunit(self, family, nparagraphs, filedb, cc_txt_file, outputdirectorychecklist, docx, xlsx):

        if not (os.path.isfile(filedb)):
          print os.linesep + "Error loading database. Must run CCParser.py first" + os.linesep
          sys.exit(1)
        # Database
        try:
            # Create or open a SQLite3 DB
            dbmgr1 = DatabaseManager(filedb)
            dbmgr2 = DatabaseManager(filedb)
            dbmgr3 = DatabaseManager(filedb)
            dbmgr4 = DatabaseManager(filedb)
            dbmgr5 = DatabaseManager(filedb)
            dbmgr6 = DatabaseManager(filedb)
            dbmgr7 = DatabaseManager(filedb)
            dbmgr8 = DatabaseManager(filedb)
            dbmgr9 = DatabaseManager(filedb)
        except NameError:
            print os.linesep + "Error loading database" + os.linesep
            sys.exit(1)

        # Initialization
        n=1 # Workunit number
        paragraph = int(nparagraphs) # number of paragraphs to write
        if (paragraph > 4): paragraph = 4
        data = dict()

        # Sub Activity
        query1 = "select child3.idelement, child3.name, child3.paratext, child5.paratext from child3, child4, child5 where child3.idelement like \"%"+family+"%\" and child3.id = child4.parentkey and child4.element=\"msa-objectives\" and child4.id = child5.parentkey group by child3.paratext"

        for row in dbmgr1.query(query1):
          txtsubactivity = row[2].upper() +": " + row[1]
          print ((colored("[-] "+ txtsubactivity, 'green')))
          printeval = "Evaluation of sub-activity (" + row[1] +": " + row[2].upper() +") " + row[3]

          print_title_txt(printeval, cc_txt_file, outputdirectorychecklist)

          docx.add_heading(row[2].upper(), level=1)
          p = docx.add_paragraph("Evaluation of sub-activity (")
          p.add_run(row[1] +": " + row[2].upper()).bold = True
          p.add_run(") " + row[3])

          xlsx.write(0, 0, txtsubactivity)


        # Workunits
        query2 = "select child4.id, child4.paratext, child4.element, child4.idelement from child3, child4 where child3.idelement like \"%"+family+"%\" and child3.id = child4.parentkey ORDER BY child3.id"
        for row in dbmgr2.query(query2):
          if (row[2]=="ae-evaluator"): 
            print_result_linesep_txt(row[1], cc_txt_file, outputdirectorychecklist)
            docx.add_paragraph(row[1])

          if (row[2]=="ae-content"): 
            data[row[3]] = row[1]

          query3 = "select child5.id, child5.paratext from child5  where  child5.element = \"m-workunit\" and child5.parentkey = "+str(row[0])+" ORDER BY child5.id"
          for row2 in dbmgr3.query(query3):
              mworkunit = family + "-" + str(n)

              print_result_txt(os.linesep, cc_txt_file, outputdirectorychecklist)
              print_result_txt("Workunit: " + mworkunit, cc_txt_file, outputdirectorychecklist)

              p2 = docx.add_paragraph()
              p2.add_run("Workunit: " + mworkunit).bold = True

              xlsx.write(n+1, 0, "Workunit: " + mworkunit)

              n = n+1

              # Workunit text
              query4 = "select child6.element, child6.paratext, child6.idelement, child6.id from child6  where child6.parentkey = "+str(row2[0])+" ORDER BY child6.id"
              n_para = 0 # paragraphs number

              for row3 in dbmgr4.query(query4):
                aedcelement = ""
                if (row3[2]):
                  aedcelement = data[row3[2]]

                  print_result_txt(aedcelement, cc_txt_file, outputdirectorychecklist)

                  docx.add_paragraph(aedcelement)

                if (row3[3] in (4009, 4068, 4141, 4209, 2779, 2841, 3948, 4023, 2801, 2868, 4083, 2825, 4156, 2895, 3025, 2674, 2750, 2677, 2753, 3477, 3495, 3571, 3695, 2564, 3016)): # AVA_VAN.1-11, AVA_VAN.1-12, etc
                  n_para = n_para +1

                  print_result_txt(row3[1], cc_txt_file, outputdirectorychecklist)

                  docx.add_paragraph(row3[1])

                  child7_para_query = 'SELECT id, parentkey, idelement, title, paratext, element FROM child7 WHERE parentkey = '+str(row3[3])+' ORDER BY id'
                  for para_item_7 in dbmgr5.query(child7_para_query):
                    if ((para_item_7[5]) == "list"):
                        child8_para_query = 'SELECT id, parentkey, idelement, title, paratext, element FROM child8 WHERE parentkey = '+str(para_item_7[0])+' ORDER BY id'
                        for para_item_8 in dbmgr6.query(child8_para_query):
                            if ((para_item_8[5]) == "item"): # element
                              #print "- " + para_item_8[4]
                              print_result_txt(("- " + para_item_8[4]), cc_txt_file, outputdirectorychecklist)
                              docx.add_paragraph(("- " + para_item_8[4]))
                              child9_para_query = 'SELECT id, parentkey, idelement, title, paratext, element FROM child9 WHERE parentkey = '+str(para_item_8[0])+' ORDER BY id'
                              for para_row_9 in dbmgr7.query(child9_para_query):
                                  if ((para_row_9[5]) == "italic" or (para_row_9[5]) == "bold" or (para_row_9[5]) == "xref"): # element
                                      print_result_txt((para_row_9[4]).encode('utf-8'), cc_txt_file, outputdirectorychecklist)
                                      docx.add_paragraph((para_row_9[4]).encode('utf-8'))

                elif (row3[3] in (1940, 3998, 2842, 1986, 3936, 2036, 2087, 2656, 2659, 2149, 2221, 4164, 2717, 2270, 2714)):
                  n_para = n_para +1
                  #print row3[1]
                  print_result_txt(row3[1], cc_txt_file, outputdirectorychecklist)
                  docx.add_paragraph(row3[1])

                  child7_para_query = 'SELECT id, parentkey, idelement, title, paratext, element FROM child7 WHERE parentkey = '+str(row3[3])+' ORDER BY id'
                  for para_item_7 in dbmgr5.query(child7_para_query):
                    if ((para_item_7[4])): # element
                        print_result_txt(para_item_7[4], cc_txt_file, outputdirectorychecklist)
                        docx.add_paragraph(para_item_7[4])
                  
                elif (n_para < paragraph):
                  num_string = len(row3[1])
                  if (row3[1] != aedcelement[:num_string]):
                    txtparagraph = row3[1]
                    child7_para_query = 'SELECT id, parentkey, idelement, title, paratext, element FROM child7 WHERE parentkey = '+str(row3[3])+' ORDER BY id'
                    for para_row_7 in dbmgr8.query(child7_para_query):
                      if ((para_row_7[5]) == "bold"):
                          txtparagraph += "<b> " + (para_row_7[4]).encode('utf-8') + " </b>"
                      if ((para_row_7[5]) == "italic"):
                          txtparagraph += "<i> " + (para_row_7[4]).encode('utf-8') + " </i>"
                      if ((para_row_7[5]) == "xref"):
                          txtparagraph += (para_row_7[4]).encode('utf-8') # paratext

                    print_result_txt(txtparagraph, cc_txt_file, outputdirectorychecklist)
                    docx.add_paragraph(txtparagraph)
                  n_para = n_para +1



###############################################################################
#------------------------------------------------------------------------------
# Checklist
#------------------------------------------------------------------------------
###############################################################################

class Checklist:

    try:
        parser = cmdline_parser()
        # Show help if no args
        if len(sys.argv) < 3:
            parser.print_help()
            sys.exit(1)
    except NameError:
        parser.print_help()
        sys.exit(1)

    # Get args
    results = parser.parse_args()
    # No failure by default :)
    failure = False

    # Database
    filedb = results.database
    # Arguments
    ealargument = results.EAL
    acomponentlist = results.list
    ealargumentprint = ""
    if (results.ac):
      ealargumentprint = ealargument + "+"
      for e in results.ac:
        ealargumentprint = ealargumentprint + " " + e
    else:
      ealargumentprint = ealargument

    if (results.npara > 0): nparagraphs = results.npara
    else: nparagraphs = 1

    # Create output directory for results
    outputdirectory = 'output'
    outputdirectorychecklist = (outputdirectory + '/checklist')
    if (os.path.isdir(outputdirectorychecklist)):
      shutil.rmtree(outputdirectorychecklist)
    if not (os.path.isdir(outputdirectory)): os.makedirs(outputdirectory)
    if not (os.path.isdir(outputdirectorychecklist)): os.makedirs(outputdirectorychecklist)
    

    cc_txt_file = 'workunit_checklist.txt'
    cc_doc_file = 'workunit_checklist.docx'
    cc_xls_file = 'workunit_checklist.xlsx'

    # Init Class
    WU = Workunit()

    A_C_EAL1 = ["ASE_CCL.1", "ASE_ECD.1", "ASE_INT.1", "ASE_OBJ.1", "ASE_REQ.1", "ASE_TSS.1", "AGD_OPE.1", "AGD_PRE.1", "ALC_CMC.1", "ALC_CMS.1", "ADV_FSP.1", "ATE_IND.1", "AVA_VAN.1"]
    A_C_EAL2 = ["ASE_CCL.1", "ASE_ECD.1", "ASE_INT.1", "ASE_OBJ.2", "ASE_REQ.2", "ASE_TSS.1", "ASE_SPD.1", "AGD_OPE.1", "AGD_PRE.1", "ALC_CMC.2", "ALC_CMS.2", "ALC_DEL.1", "ADV_ARC.1", "ADV_FSP.2", "ADV_TDS.1", "ATE_COV.1", "ATE_FUN.1", "ATE_IND.2", "AVA_VAN.2"]
    A_C_EAL3 = ["ASE_CCL.1", "ASE_ECD.1", "ASE_INT.1", "ASE_OBJ.2", "ASE_REQ.2", "ASE_TSS.1", "ASE_SPD.1", "AGD_OPE.1", "AGD_PRE.1", "ALC_CMC.3", "ALC_CMS.3", "ALC_DEL.1", "ALC_DVS.1", "ALC_LCD.1", "ADV_ARC.1", "ADV_FSP.3", "ADV_TDS.2", "ATE_COV.2", "ATE_DPT.1", "ATE_FUN.1", "ATE_IND.2", "AVA_VAN.2"]
    A_C_EAL4 = ["ASE_CCL.1", "ASE_ECD.1", "ASE_INT.1", "ASE_OBJ.2", "ASE_REQ.2", "ASE_TSS.1", "ASE_SPD.1", "AGD_OPE.1", "AGD_PRE.1", "ALC_CMC.4", "ALC_CMS.4", "ALC_DEL.1", "ALC_DVS.1", "ALC_LCD.1", "ALC_TAT.1", "ADV_ARC.1", "ADV_FSP.4", "ADV_TDS.3", "ADV_IMP.1", "ATE_COV.2", "ATE_DPT.1", "ATE_FUN.1", "ATE_IND.2", "AVA_VAN.3"]
    A_C_EAL5 = ["ASE_CCL.1", "ASE_ECD.1", "ASE_INT.1", "ASE_OBJ.2", "ASE_REQ.2", "ASE_TSS.1", "ASE_SPD.1", "AGD_OPE.1", "AGD_PRE.1", "ALC_CMC.4", "ALC_CMS.5", "ALC_DEL.1", "ALC_DVS.1", "ALC_LCD.1", "ALC_TAT.2", "ADV_ARC.1", "ADV_FSP.5", "ADV_TDS.4", "ADV_IMP.1", "ADV_INT.2", "ATE_COV.2", "ATE_DPT.3", "ATE_FUN.1", "ATE_IND.2", "AVA_VAN.4"]
    A_C_EAL6 = ["ASE_CCL.1", "ASE_ECD.1", "ASE_INT.1", "ASE_OBJ.2", "ASE_REQ.2", "ASE_TSS.1", "ASE_SPD.1", "AGD_OPE.1", "AGD_PRE.1", "ALC_CMC.5", "ALC_CMS.5", "ALC_DEL.1", "ALC_DVS.2", "ALC_LCD.1", "ALC_TAT.3", "ADV_ARC.1", "ADV_FSP.5", "ADV_TDS.5", "ADV_IMP.2", "ADV_INT.3", "ADV_SPM.1", "ATE_COV.3", "ATE_DPT.3", "ATE_FUN.2", "ATE_IND.2", "AVA_VAN.5"]
    A_C_EAL7 = ["ASE_CCL.1", "ASE_ECD.1", "ASE_INT.1", "ASE_OBJ.2", "ASE_REQ.2", "ASE_TSS.1", "ASE_SPD.1", "AGD_OPE.1", "AGD_PRE.1", "ALC_CMC.5", "ALC_CMS.5", "ALC_DEL.1", "ALC_DVS.2", "ALC_LCD.2", "ALC_TAT.3", "ADV_ARC.1", "ADV_FSP.6", "ADV_TDS.6", "ADV_IMP.2", "ADV_INT.3", "ADV_SPM.1", "ATE_COV.3", "ATE_DPT.4", "ATE_FUN.2", "ATE_IND.3", "AVA_VAN.5"]


    # ----------------------------------------------------------------------------------------------------------------#

    if (acomponentlist):
      WU.AugmentationList(filedb)
    # ----------------------------------------------------------------------------------------------------------------#
    # Create the results files

    create_txt_file(cc_txt_file, outputdirectorychecklist)
    document = Document()

    # ----------------------------------------------------------------------------------------------------------------#
    # Create an new Excel file and add a worksheet.
    workbook = xlsxwriter.Workbook(outputdirectorychecklist+"/"+cc_xls_file)
    bold = workbook.add_format({'bold': True})
    italic = workbook.add_format({'italic': True})
    red = workbook.add_format({'color': 'red'})
    blue = workbook.add_format({'color': 'blue'})
    center = workbook.add_format({'align': 'center'})
    superscript = workbook.add_format({'font_script': 1})

    cell_format = workbook.add_format({'align': 'center',
                                       'valign': 'vcenter',
                                       'border': 1})

    # ----------------------------------------------------------------------------------------------------------------#

    ealtxt = "Evaluation Assurance Level: " + ealargumentprint
    print ((colored("[+] "+ ealtxt, 'red')))
    print_title_txt(ealtxt, cc_txt_file, outputdirectorychecklist)
    document.add_heading(ealtxt, 0)

    worksheet = workbook.add_worksheet(ealargument)
    worksheet.merge_range('C2:I10', "", cell_format)
    worksheet.write_rich_string('C2', red, ealargumentprint, cell_format)

    A_C_EAL_SELECTED = ""
    if (ealargument == "EAL1"):
      A_C_EAL_SELECTED = A_C_EAL1
      for assurance_family in A_C_EAL1:
        worksheet = workbook.add_worksheet(assurance_family)
        worksheet.merge_range('A1:F1', "", cell_format)
        worksheet.set_column('A:A', 30)
        WU.queryworkunit(assurance_family, nparagraphs, filedb, cc_txt_file, outputdirectorychecklist, document, worksheet)
    elif (ealargument == "EAL2"):
      A_C_EAL_SELECTED = A_C_EAL2
      for assurance_family in A_C_EAL2:
        worksheet = workbook.add_worksheet(assurance_family)
        worksheet.merge_range('A1:F1', "", cell_format)
        worksheet.set_column('A:A', 30)
        WU.queryworkunit(assurance_family, nparagraphs, filedb, cc_txt_file, outputdirectorychecklist, document, worksheet)
    elif (ealargument == "EAL3"):
      A_C_EAL_SELECTED = A_C_EAL3
      for assurance_family in A_C_EAL3:
        worksheet = workbook.add_worksheet(assurance_family)
        worksheet.merge_range('A1:F1', "", cell_format)
        worksheet.set_column('A:A', 30)
        WU.queryworkunit(assurance_family, nparagraphs, filedb, cc_txt_file, outputdirectorychecklist, document, worksheet)
    elif (ealargument == "EAL4"):
      A_C_EAL_SELECTED = A_C_EAL4
      for assurance_family in A_C_EAL4:
        worksheet = workbook.add_worksheet(assurance_family)
        worksheet.merge_range('A1:F1', "", cell_format)
        worksheet.set_column('A:A', 30)
        WU.queryworkunit(assurance_family, nparagraphs, filedb, cc_txt_file, outputdirectorychecklist, document, worksheet)
    elif (ealargument == "EAL5"):
      A_C_EAL_SELECTED = A_C_EAL5
      for assurance_family in A_C_EAL5:
        worksheet = workbook.add_worksheet(assurance_family)
        worksheet.merge_range('A1:F1', "", cell_format)
        worksheet.set_column('A:A', 30)
        WU.queryworkunit(assurance_family, nparagraphs, filedb, cc_txt_file, outputdirectorychecklist, document, worksheet)
    elif (ealargument == "EAL6"):
      A_C_EAL_SELECTED = A_C_EAL6
      for assurance_family in A_C_EAL6:
        worksheet = workbook.add_worksheet(assurance_family)
        worksheet.merge_range('A1:F1', "", cell_format)
        worksheet.set_column('A:A', 30)
        WU.queryworkunit(assurance_family, nparagraphs, filedb, cc_txt_file, outputdirectorychecklist, document, worksheet)
    elif (ealargument == "EAL7"):
      A_C_EAL_SELECTED = A_C_EAL7
      for assurance_family in A_C_EAL7:
        worksheet = workbook.add_worksheet(assurance_family)
        worksheet.merge_range('A1:F1', "", cell_format)
        worksheet.set_column('A:A', 30)
        WU.queryworkunit(assurance_family, nparagraphs, filedb, cc_txt_file, outputdirectorychecklist, document, worksheet)

    # ----------------------------------------------------------------------------------------------------------------#
    if (results.ac):
      for acomponent in results.ac:
        if (acomponent not in A_C_EAL_SELECTED):
          worksheet = workbook.add_worksheet(acomponent)
          worksheet.merge_range('A1:F1', "", cell_format)
          worksheet.set_column('A:A', 30)
          WU.queryworkunit(acomponent, nparagraphs, filedb, cc_txt_file, outputdirectorychecklist, document, worksheet)

    # ----------------------------------------------------------------------------------------------------------------#

    document.save(outputdirectorychecklist+"/"+cc_doc_file)
    workbook.close()

    # Press ENTER to continue
    if (results.ask):
        python2 = sys.version_info[0] == 2
        if python2:
            raw_input("Press ENTER to continue")
        else:
            input("Press ENTER to continue")

#------------------------------------------------------------------------------
# Main
#------------------------------------------------------------------------------

if __name__ == "__main__":
    if Checklist.failure:
        sys.exit(1)
